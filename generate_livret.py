"""
Génère le livret de présentation Word (.docx) pour BININGA
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs ────────────────────────────────────────────────────────────────
ROUGE   = RGBColor(0xC0, 0x00, 0x00)   # rouge institutionnel
BLANC   = RGBColor(0xFF, 0xFF, 0xFF)
NOIR    = RGBColor(0x1A, 0x1A, 0x1A)
GRIS    = RGBColor(0x55, 0x55, 0x55)
GRIS_C  = RGBColor(0xF5, 0xF5, 0xF5)
OR      = RGBColor(0xB8, 0x96, 0x00)

# ── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="C00000", size=6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def para_space(doc, before=0, after=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    return p

def heading(doc, text, level=1, color=ROUGE, size=20, bold=True, center=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after  = Pt(4)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def body(doc, text, color=NOIR, size=11, italic=False, center=False, bold=False, before=0, after=4):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    run.bold = bold
    return p

def divider(doc, color="C00000"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p

# ── Document ─────────────────────────────────────────────────────────────────
doc = Document()

# Marges
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Police par défaut
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE COUVERTURE
# ══════════════════════════════════════════════════════════════════════════════

# Bandeau rouge haut
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t.cell(0,0)
set_cell_bg(cell, "C00000")
cell.width = Cm(16)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(14)
run = p.add_run("PRÉSENTATION OFFICIELLE")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = BLANC
run.font.name = 'Calibri'
no_borders(t)

para_space(doc, after=16)

# Nom principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("Ange Aimé Wilfrid")
run.bold = False
run.font.size = Pt(22)
run.font.color.rgb = GRIS
run.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(20)
run2 = p2.add_run("BININGA")
run2.bold = True
run2.font.size = Pt(42)
run2.font.color.rgb = ROUGE
run2.font.name = 'Calibri'

divider(doc)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(8)
p3.paragraph_format.space_after  = Pt(6)
for line in [
    "Garde des Sceaux · Ministre de la Justice",
    "des Droits Humains et de la Promotion des Peuples Autochtones",
    "Député de la 1re Circonscription d'Ewo · PCT",
]:
    run = p3.add_run(line + "\n")
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS
    run.font.name = 'Calibri'

para_space(doc, after=30)

# Bloc site officiel
t2 = doc.add_table(rows=1, cols=1)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
c2 = t2.cell(0,0)
set_cell_bg(c2, "F5F5F5")
set_cell_borders(c2, "C00000", 12)
p4 = c2.paragraphs[0]
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(14)
p4.paragraph_format.space_after  = Pt(14)
r1 = p4.add_run("Site Web Officiel — Dossier de Présentation Technique\n")
r1.bold = True
r1.font.size = Pt(13)
r1.font.color.rgb = ROUGE
r1.font.name = 'Calibri'
r2 = p4.add_run("Développé et livré par :")
r2.font.size = Pt(10)
r2.font.color.rgb = GRIS
r2.font.name = 'Calibri'
no_borders(t2)

para_space(doc, after=40)

# Date
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p5.add_run("Mars 2026  ·  République du Congo")
r.font.size = Pt(10)
r.font.color.rgb = GRIS
r.italic = True
r.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "1.  Contexte et Objectif", size=16, color=ROUGE)
divider(doc)

body(doc, (
    "Dans un monde où la communication numérique est devenue un pilier incontournable de "
    "l'action politique, il est essentiel pour un homme d'État de disposer d'une présence "
    "en ligne à la hauteur de son engagement et de sa stature."
), size=11, before=4, after=6)

body(doc, (
    "C'est dans cette optique qu'a été conçu et développé le site officiel de "
    "l'Honorable Ange Aimé Wilfrid BININGA, Garde des Sceaux, Ministre de la Justice "
    "et Député d'Ewo — un outil de communication institutionnel, moderne et entièrement "
    "personnalisable, pensé pour servir son action au quotidien."
), size=11, before=0, after=10)

# Encadré slogan
t3 = doc.add_table(rows=1, cols=1)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
c3 = t3.cell(0,0)
set_cell_bg(c3, "C00000")
p6 = c3.paragraphs[0]
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
p6.paragraph_format.space_before = Pt(12)
p6.paragraph_format.space_after  = Pt(12)
r = p6.add_run('"Pour Ewo. Pour un Congo juste, libre et fort."')
r.bold = True
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = BLANC
r.font.name = 'Calibri'
no_borders(t3)

para_space(doc, after=12)

body(doc, (
    "Le site a été entièrement développé sur mesure — sans template, sans CMS générique, "
    "sans code préfabriqué. Chaque ligne de code a été écrite spécifiquement pour "
    "répondre aux besoins de communication d'un ministre et député congolais."
), size=11, before=4, after=6)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — CE QUI A ÉTÉ LIVRÉ
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "2.  Ce Qui a Été Livré", size=16, color=ROUGE)
divider(doc)

body(doc, "Le projet comprend trois composantes distinctes et complémentaires :", size=11, before=4, after=8)

# Tableau des 3 composantes
t4 = doc.add_table(rows=4, cols=3)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(t4)

headers = ["Composante", "Description", "Complexité"]
for i, h in enumerate(headers):
    c = t4.cell(0, i)
    set_cell_bg(c, "1A1A1A")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = BLANC
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

rows_data = [
    ("Site public officiel", "Vitrine web complète du ministre, 8 sections, design premium, 100% dynamique", "Élevée"),
    ("Panneau d'administration", "Interface de gestion complète, 10 modules, gestion des messages citoyens", "Très élevée"),
    ("Serveur & API sécurisée", "Serveur Python sur mesure, authentification, sauvegarde automatique, HTTPS", "Élevée"),
]
bg_alt = ["FAFAFA", "F0F0F0", "FAFAFA"]
for row_i, (comp, desc, compl) in enumerate(rows_data):
    row = t4.rows[row_i+1]
    bg = bg_alt[row_i]
    for ci, val in enumerate([comp, desc, compl]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after  = Pt(7)
        if ci == 0:
            r = p.add_run(val)
            r.bold = True
            r.font.color.rgb = ROUGE
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
        elif ci == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.bold = True
            r.font.color.rgb = NOIR
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
        else:
            r = p.add_run(val)
            r.font.color.rgb = GRIS
            r.font.size = Pt(10)
            r.font.name = 'Calibri'

para_space(doc, after=16)
heading(doc, "Volume de travail réalisé", size=13, color=NOIR, bold=True)

# Stats lignes de code
t5 = doc.add_table(rows=5, cols=3)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(t5)
stats_rows = [
    ("Fichier", "Lignes de code", "Poids"),
    ("index.html  (site public)", "1 257 lignes", "78 Ko"),
    ("admin.html  (panneau admin)", "1 096 lignes", "53 Ko"),
    ("server.py  (serveur)", "230 lignes", "10 Ko"),
    ("TOTAL", "2 725 lignes", "146 Ko"),
]
for ri, (a, b, c) in enumerate(stats_rows):
    row = t5.rows[ri]
    is_header = ri == 0
    is_total  = ri == 4
    bg = "C00000" if is_header else ("1A1A1A" if is_total else ("F5F5F5" if ri%2==0 else "EBEBEB"))
    fg = BLANC if (is_header or is_total) else NOIR
    for ci, val in enumerate([a, b, c]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after  = Pt(7)
        if ci != 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        r.bold = is_header or is_total
        r.font.color.rgb = fg
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — FONCTIONNALITÉS SITE PUBLIC
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "3.  Fonctionnalités — Site Public", size=16, color=ROUGE)
divider(doc)
body(doc, "Le site public comprend 8 sections, toutes accessibles en un clic depuis la navigation :", size=11, before=4, after=8)

sections = [
    ("ACCUEIL (Hero)", "Animation d'entrée cinématique, nom animé, slogan politique, appel à l'action."),
    ("À PROPOS", "Biographie du ministre, présentation de son parcours, 3 compteurs animés (mandats, ministères, circonscription)."),
    ("ENGAGEMENT", "Déclaration des valeurs et engagements politiques du député."),
    ("PARCOURS", "Ligne du temps interactive retraçant les étapes clés de sa carrière."),
    ("PROGRAMME", "Les priorités et le programme politique pour Ewo et le Congo."),
    ("GALERIE", "Slider photos avec légendes + grille photos cliquable, 100% gérable depuis l'admin."),
    ("ACTUALITÉS", "Article principal mis en avant + liste des dernières actualités, gérée depuis l'admin."),
    ("CONTACT & AUDIENCES", "Formulaire de demande d'audience officielle + formulaire de contact, avec envoi email réel."),
]

t6 = doc.add_table(rows=len(sections)+1, cols=2)
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(t6)

# Header
for ci, h in enumerate(["Section", "Contenu"]):
    c = t6.cell(0, ci)
    set_cell_bg(c, "C00000")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = BLANC
    r.font.size = Pt(10)
    r.font.name = 'Calibri'

for ri, (sec, desc) in enumerate(sections):
    bg = "F5F5F5" if ri%2==0 else "EBEBEB"
    row = t6.rows[ri+1]
    c0 = row.cells[0]
    set_cell_bg(c0, bg)
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after  = Pt(6)
    r0 = p0.add_run(sec)
    r0.bold = True
    r0.font.color.rgb = ROUGE
    r0.font.size = Pt(10)
    r0.font.name = 'Calibri'

    c1 = row.cells[1]
    set_cell_bg(c1, bg)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(6)
    r1 = p1.add_run(desc)
    r1.font.color.rgb = GRIS
    r1.font.size = Pt(10)
    r1.font.name = 'Calibri'

para_space(doc, after=12)
body(doc, "Points techniques notables :", bold=True, size=11, before=4, after=4)

tech_points = [
    "Curseur rouge animé personnalisé — expérience utilisateur premium",
    "Loader animé à l'ouverture de la page — impression de fluidité professionnelle",
    "Animations au défilement (IntersectionObserver) — chaque section s'anime à l'entrée",
    "100% du contenu est dynamique — aucun texte codé en dur, tout modifiable depuis l'admin",
    "Protection contre les injections de code malveillant (anti-XSS)",
    "Responsive complet — s'adapte parfaitement sur téléphone, tablette et ordinateur",
    "Menu burger mobile élégant",
    "Envoi email réel via Formspree (audiences et contacts)",
    "Compteur de visiteurs intégré",
]
for pt in tech_points:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run("✔  " + pt)
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS
    run.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — PANNEAU ADMINISTRATION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "4.  Panneau d'Administration — Le CMS", size=16, color=ROUGE)
divider(doc)

body(doc, (
    "Le panneau d'administration est un système de gestion de contenu (CMS) développé "
    "entièrement sur mesure. Il permet à Son Excellence ou à un collaborateur désigné "
    "de gérer l'intégralité du site sans aucune compétence technique."
), size=11, before=4, after=8)

admin_panels = [
    ("Tableau de bord", "Vue d'ensemble en temps réel : nombre de visiteurs, demandes d'audience en attente, réclamations non traitées, messages non lus."),
    ("Éditeur Hero", "Modifier le nom affiché, le rôle officiel, le slogan politique et le sous-titre de la page d'accueil."),
    ("Éditeur À propos", "Modifier le titre de section et l'introduction biographique."),
    ("Éditeur Statistiques", "Mettre à jour les 3 compteurs (mandats, ministères, circonscription)."),
    ("Éditeur SEO", "Contrôler le titre et la description qui apparaissent dans Google."),
    ("Éditeur Galerie Slider", "Ajouter, supprimer, réordonner les photos du slider. Uploader de nouvelles images directement."),
    ("Éditeur Galerie Grille", "Gérer les photos de la grille galerie, ajout et suppression en un clic."),
    ("Éditeur Actualités", "Modifier l'article principal mis en avant + gérer la liste des actualités (date, titre, description)."),
    ("Demandes d'audiences", "Consulter, filtrer et traiter toutes les demandes d'audience soumises par les citoyens."),
    ("Réclamations", "Consulter et gérer toutes les réclamations des citoyens avec statut (nouveau/en cours/traité)."),
    ("Messages contact", "Consulter tous les messages de contact avec suivi du statut de traitement."),
]

t7 = doc.add_table(rows=len(admin_panels)+1, cols=2)
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(t7)

for ci, h in enumerate(["Module Admin", "Fonctionnalité"]):
    c = t7.cell(0, ci)
    set_cell_bg(c, "1A1A1A")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(7)
    if ci==1: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = BLANC
    r.font.size = Pt(10)
    r.font.name = 'Calibri'

for ri, (mod, feat) in enumerate(admin_panels):
    bg = "F5F5F5" if ri%2==0 else "EBEBEB"
    row = t7.rows[ri+1]
    c0 = row.cells[0]
    set_cell_bg(c0, bg)
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    r0 = p0.add_run(mod)
    r0.bold = True
    r0.font.color.rgb = NOIR
    r0.font.size = Pt(10)
    r0.font.name = 'Calibri'

    c1 = row.cells[1]
    set_cell_bg(c1, bg)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(5)
    r1 = p1.add_run(feat)
    r1.font.color.rgb = GRIS
    r1.font.size = Pt(10)
    r1.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — SÉCURITÉ & TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "5.  Sécurité & Architecture Technique", size=16, color=ROUGE)
divider(doc)

body(doc, "Le site a été conçu avec un niveau de sécurité adapté à un site institutionnel :", size=11, before=4, after=8)

sec_items = [
    ("Authentification serveur", "Le mot de passe n'est jamais stocké en clair. La vérification se fait côté serveur Python."),
    ("Token de session dynamique", "Chaque session génère un token unique. Pas de connexion admin possible sans ce token."),
    ("Protection HTTPS", "Le serveur supporte le chiffrement SSL/TLS natif — connexion sécurisée de bout en bout."),
    ("Sauvegarde automatique", "À chaque modification, une copie de sauvegarde du contenu est créée automatiquement."),
    ("Anti-XSS", "Toutes les données saisies par les utilisateurs sont neutralisées avant affichage (protection contre les injections de code)."),
    ("Limitation des uploads", "Seules les images (JPG, PNG, WebP, GIF) sont acceptées. Taille maximum contrôlée."),
    ("Isolation admin", "Le panneau admin est sur une URL séparée, protégée par login — invisible aux visiteurs."),
]

for title, desc in sec_items:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run("► " + title + "  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = ROUGE
    r1.font.name = 'Calibri'
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRIS
    r2.font.name = 'Calibri'

para_space(doc, after=14)
heading(doc, "Technologies utilisées", size=13, color=NOIR, bold=True)
body(doc, "Le projet repose exclusivement sur des technologies open-source, sans coût de licence :", size=10, before=4, after=6)

tech_table = [
    ("HTML5 / CSS3", "Structure et design du site"),
    ("JavaScript (vanilla)", "Toute la logique interactive — sans framework (plus léger, plus rapide)"),
    ("Python 3", "Serveur web et API — aucune dépendance externe"),
    ("JSON", "Base de données légère pour le contenu du site"),
    ("Formspree", "Service d'envoi d'emails (plan gratuit disponible)"),
    ("Google Fonts", "Typographie premium (Raleway, Lato)"),
]
t8 = doc.add_table(rows=len(tech_table)+1, cols=2)
no_borders(t8)
for ci, h in enumerate(["Technologie", "Usage"]):
    c = t8.cell(0, ci)
    set_cell_bg(c, "C00000")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = BLANC
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
for ri, (tech, use) in enumerate(tech_table):
    bg = "F5F5F5" if ri%2==0 else "EBEBEB"
    row = t8.rows[ri+1]
    for ci, val in enumerate([tech, use]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(val)
        r.bold = ci==0
        r.font.color.rgb = ROUGE if ci==0 else GRIS
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 7 — VALEUR & ESTIMATION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "6.  Valeur du Projet", size=16, color=ROUGE)
divider(doc)

body(doc, (
    "Ce projet a été réalisé en solo, par un développeur unique ayant endossé simultanément "
    "les rôles de designer, développeur front-end, développeur back-end, expert en sécurité "
    "et chef de projet. Dans une agence professionnelle, ce travail aurait nécessité :"
), size=11, before=4, after=8)

team = [
    ("Designer UI/UX", "Maquettes, identité visuelle, animations, typographie"),
    ("Intégrateur Front-End", "HTML, CSS, responsive design, slider, animations"),
    ("Développeur JavaScript", "Logique dynamique, CMS admin, formulaires"),
    ("Développeur Back-End", "Serveur Python, API REST, authentification, sécurité"),
    ("Chef de Projet", "Coordination, livraison, suivi client"),
]

t9 = doc.add_table(rows=len(team)+1, cols=2)
no_borders(t9)
for ci, h in enumerate(["Rôle", "Mission"]):
    c = t9.cell(0, ci)
    set_cell_bg(c, "1A1A1A")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(7)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = BLANC
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
for ri, (role, mission) in enumerate(team):
    bg = "F5F5F5" if ri%2==0 else "EBEBEB"
    row = t9.rows[ri+1]
    for ci, val in enumerate([role, mission]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(val)
        r.bold = ci==0
        r.font.color.rgb = ROUGE if ci==0 else GRIS
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

para_space(doc, after=14)
heading(doc, "Estimation comparative du coût", size=13, color=NOIR, bold=True)

pricing = [
    ("Freelance Congo / Afrique Centrale", "400 000 – 800 000 FCFA"),
    ("Agence web africaine", "800 000 – 2 000 000 FCFA"),
    ("Agence web France / Europe", "4 000 – 10 000 €"),
]

t10 = doc.add_table(rows=len(pricing)+1, cols=2)
no_borders(t10)
for ci, h in enumerate(["Contexte de référence", "Estimation de marché"]):
    c = t10.cell(0, ci)
    set_cell_bg(c, "C00000")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(7)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = BLANC
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
for ri, (ctx, price) in enumerate(pricing):
    bg = "F5F5F5" if ri%2==0 else "EBEBEB"
    row = t10.rows[ri+1]
    for ci, val in enumerate([ctx, price]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(val)
        r.bold = ci==1
        r.font.color.rgb = GRIS if ci==0 else ROUGE
        r.font.size = Pt(11 if ci==1 else 10)
        r.font.name = 'Calibri'

para_space(doc, after=10)
body(doc, (
    "La valeur ajoutée réelle de ce projet réside dans le CMS sur mesure (panneau admin), "
    "qui à lui seul représente 60% du coût de développement dans tout devis professionnel."
), size=10, italic=True, color=GRIS, before=4, after=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 8 — MAINTENANCE & PROPOSITION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "7.  Maintenance & Continuité", size=16, color=ROUGE)
divider(doc)

body(doc, (
    "Un site web n'est pas un projet ponctuel — c'est un outil vivant qui nécessite "
    "un suivi technique régulier pour rester performant, sécurisé et à jour."
), size=11, before=4, after=8)

maintenance = [
    ("Mises à jour de sécurité", "Vérification et application des correctifs de sécurité du serveur."),
    ("Évolutions du contenu", "Ajout de nouvelles fonctionnalités selon les besoins (nouvelles sections, nouveaux modules)."),
    ("Sauvegarde des données", "Sauvegarde régulière des données et des photos du site."),
    ("Hébergement serveur", "Gestion du serveur et du nom de domaine officiel."),
    ("Support technique", "Assistance en cas de problème ou de panne."),
    ("Formation équipe", "Formation des collaborateurs à l'utilisation du panneau admin."),
]

for title, desc in maintenance:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run("■  " + title + " : ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = ROUGE
    r1.font.name = 'Calibri'
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRIS
    r2.font.name = 'Calibri'

para_space(doc, after=20)

# Bloc conclusion
t11 = doc.add_table(rows=1, cols=1)
t11.alignment = WD_TABLE_ALIGNMENT.CENTER
c11 = t11.cell(0,0)
set_cell_bg(c11, "F5F5F5")
set_cell_borders(c11, "C00000", 12)
p11 = c11.paragraphs[0]
p11.alignment = WD_ALIGN_PARAGRAPH.CENTER
p11.paragraph_format.space_before = Pt(14)
p11.paragraph_format.space_after  = Pt(14)
r = p11.add_run(
    "Ce site officiel est prêt à être déployé et présenté au public.\n"
    "Il constitue un outil de communication institutionnelle moderne,\n"
    "digne de la stature de Son Excellence BININGA."
)
r.font.size = Pt(12)
r.font.color.rgb = NOIR
r.font.name = 'Calibri'
no_borders(t11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 9 — CONTACT DÉVELOPPEUR
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "8.  Le Développeur", size=16, color=ROUGE)
divider(doc)

body(doc, (
    "Ce projet a été intégralement conçu, designé et développé par un développeur "
    "full-stack congolais, maîtrisant à la fois le design, le développement front-end, "
    "le développement back-end et la sécurité informatique."
), size=11, before=4, after=8)

skills = [
    "Design UI/UX et identité visuelle",
    "Développement front-end (HTML, CSS, JavaScript)",
    "Développement back-end (Python, API REST)",
    "Sécurité informatique (authentification, chiffrement, anti-XSS)",
    "Gestion de projet et livraison client",
    "Formation et transfert de compétences",
]

body(doc, "Compétences couvertes sur ce projet :", bold=True, size=11, before=0, after=6)
for sk in skills:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("✔  " + sk)
    r.font.size = Pt(10)
    r.font.color.rgb = GRIS
    r.font.name = 'Calibri'

para_space(doc, after=20)

# Bandeau rouge final
t_final = doc.add_table(rows=1, cols=1)
t_final.alignment = WD_TABLE_ALIGNMENT.CENTER
cf = t_final.cell(0,0)
set_cell_bg(cf, "C00000")
pf2 = cf.paragraphs[0]
pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf2.paragraph_format.space_before = Pt(18)
pf2.paragraph_format.space_after  = Pt(18)
r1 = pf2.add_run("République du Congo  ·  Mars 2026\n")
r1.font.size = Pt(11)
r1.font.color.rgb = BLANC
r1.font.name = 'Calibri'
r2 = pf2.add_run("Site officiel — Ange Aimé Wilfrid BININGA")
r2.bold = True
r2.font.size = Pt(13)
r2.font.color.rgb = BLANC
r2.font.name = 'Calibri'
no_borders(t_final)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
doc.save("Livret_Presentation_BININGA.docx")
print("✅  Livret Word généré : Livret_Presentation_BININGA.docx")
